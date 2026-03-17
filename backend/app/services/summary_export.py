"""
Export summary data to Markdown and DOCX formats.
"""
import io

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Parliament of Barbados brand colours
NAVY = RGBColor(0x00, 0x30, 0x87)
GOLD = RGBColor(0xFD, 0xB9, 0x13)


def _fmt_ts(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int(seconds % 3600 // 60)
    s = int(seconds % 60)
    return f"{h}:{m:02d}:{s:02d}" if h else f"{m}:{s:02d}"


# ── Markdown ──────────────────────────────────────────────────────────────────

def summary_to_md(summary: dict, title: str) -> str:
    lines: list[str] = []
    if title:
        lines.append(f"# {title}\n")

    lines += ["## Executive Summary\n", summary.get("executive_summary", "") + "\n"]

    topics = summary.get("topics") or []
    if topics:
        lines.append("## Topics Discussed\n")
        for topic in topics:
            ts = _fmt_ts(topic.get("start_time", 0))
            lines.append(f"### {topic.get('title', 'Unknown')}  `[{ts}]`\n")
            lines.append(topic.get("summary", "") + "\n")
            sp = topic.get("speakers") or []
            if sp:
                lines.append(f"**Speakers:** {', '.join(sp)}\n")

    decisions = summary.get("decisions") or []
    if decisions:
        lines.append("## Decisions & Votes\n")
        for d in decisions:
            ts = _fmt_ts(d.get("timestamp", 0))
            outcome = d.get("outcome", "unknown").upper()
            lines.append(f"- **{outcome}** — {d.get('description', '')}  `[{ts}]`")
        lines.append("")

    actions = summary.get("actions") or []
    if actions:
        lines.append("## Action Items\n")
        for i, a in enumerate(actions, 1):
            ts = _fmt_ts(a.get("timestamp", 0))
            resp = f"  *(— {a['responsible']})*" if a.get("responsible") else ""
            lines.append(f"{i}. {a.get('description', '')}{resp}  `[{ts}]`")
        lines.append("")

    speakers = summary.get("speakers") or []
    if speakers:
        lines.append("## Speakers\n")
        for sp in speakers:
            lines.append(f"### {sp.get('name', 'Unknown')}")
            if sp.get("role"):
                lines.append(f"*{sp['role']}*\n")
            for pos in (sp.get("key_positions") or []):
                lines.append(f"- {pos}")
            lines.append("")

    return "\n".join(lines)


# ── DOCX ──────────────────────────────────────────────────────────────────────

def _set_run_colour(run, colour: RGBColor):
    run.font.color.rgb = colour


def _add_coloured_heading(doc: Document, text: str, level: int, colour: RGBColor = NAVY):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        _set_run_colour(run, colour)
    return heading


def summary_to_docx(summary: dict, title: str) -> bytes:
    doc = Document()

    # Remove default wide margins — bring to 1 inch
    for section in doc.sections:
        section.top_margin = int(914400)       # 1 inch in EMUs
        section.bottom_margin = int(914400)
        section.left_margin = int(914400)
        section.right_margin = int(914400)

    # ── Title ──────────────────────────────────────────────────────────────────
    title_para = doc.add_heading(title or "Parliamentary Session Summary", 0)
    for run in title_para.runs:
        _set_run_colour(run, NAVY)

    # ── Executive Summary ──────────────────────────────────────────────────────
    _add_coloured_heading(doc, "Executive Summary", 1)
    doc.add_paragraph(summary.get("executive_summary", ""))

    # ── Topics ─────────────────────────────────────────────────────────────────
    topics = summary.get("topics") or []
    if topics:
        _add_coloured_heading(doc, "Topics Discussed", 1)
        for topic in topics:
            ts = _fmt_ts(topic.get("start_time", 0))
            h2 = doc.add_heading("", level=2)
            run_title = h2.add_run(topic.get("title", "Unknown"))
            _set_run_colour(run_title, NAVY)
            run_ts = h2.add_run(f"  [{ts}]")
            run_ts.font.size = Pt(9)
            run_ts.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

            doc.add_paragraph(topic.get("summary", ""))
            sp = topic.get("speakers") or []
            if sp:
                p = doc.add_paragraph()
                run_label = p.add_run("Speakers: ")
                run_label.bold = True
                p.add_run(", ".join(sp))

    # ── Decisions ──────────────────────────────────────────────────────────────
    decisions = summary.get("decisions") or []
    if decisions:
        _add_coloured_heading(doc, "Decisions & Votes", 1)
        for d in decisions:
            ts = _fmt_ts(d.get("timestamp", 0))
            outcome = d.get("outcome", "unknown").upper()
            p = doc.add_paragraph(style="List Bullet")
            badge = p.add_run(f"[{outcome}]  ")
            badge.bold = True
            # Colour the badge based on outcome
            badge.font.color.rgb = (
                RGBColor(0x00, 0x6B, 0x3F) if outcome == "PASSED"
                else RGBColor(0xCC, 0x00, 0x00) if outcome == "DEFEATED"
                else RGBColor(0xC0, 0x80, 0x00)
            )
            p.add_run(f"{d.get('description', '')}  [{ts}]")

    # ── Actions ────────────────────────────────────────────────────────────────
    actions = summary.get("actions") or []
    if actions:
        _add_coloured_heading(doc, "Action Items", 1)
        for a in actions:
            ts = _fmt_ts(a.get("timestamp", 0))
            p = doc.add_paragraph(style="List Number")
            p.add_run(a.get("description", ""))
            if a.get("responsible"):
                resp_run = p.add_run(f"  — {a['responsible']}")
                resp_run.italic = True
            grey = p.add_run(f"  [{ts}]")
            grey.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ── Speakers ───────────────────────────────────────────────────────────────
    speakers = summary.get("speakers") or []
    if speakers:
        _add_coloured_heading(doc, "Speakers", 1)
        for sp in speakers:
            _add_coloured_heading(doc, sp.get("name", "Unknown"), 2)
            if sp.get("role"):
                role_p = doc.add_paragraph()
                role_run = role_p.add_run(sp["role"])
                role_run.italic = True
                role_run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            for pos in (sp.get("key_positions") or []):
                doc.add_paragraph(pos, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
